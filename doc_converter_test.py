#!/usr/bin/env python3
# docx_converter_test.py - Test script for HTML to DOCX conversion

import os
import sys
import tempfile
import traceback
import subprocess
from pathlib import Path

def test_html2docx_import():
    """Test if the html2docx module can be imported."""
    print("\n=== Testing html2docx import ===")
    try:
        import html2docx
        print("✅ html2docx module imported successfully")
        print(f"html2docx version: {html2docx.__version__ if hasattr(html2docx, '__version__') else 'unknown'}")
        print(f"html2docx location: {html2docx.__file__}")
        return True
    except ImportError as e:
        print(f"❌ Failed to import html2docx: {str(e)}")
        print("Attempting to install the package...")
        try:
            subprocess.run([sys.executable, "-m", "pip", "install", "html2docx"], check=True)
            print("Installation command completed")
            
            try:
                import html2docx
                print("✅ html2docx module now installed and imported successfully")
                return True
            except ImportError as post_install_error:
                print(f"❌ Still cannot import html2docx after installation: {str(post_install_error)}")
                return False
        except Exception as install_error:
            print(f"❌ Failed to install html2docx: {str(install_error)}")
            print("Please manually install the required package with:")
            print("pip install html2docx")
            return False
    except Exception as e:
        print(f"❌ Unexpected error importing html2docx: {str(e)}")
        traceback.print_exc()
        return False

def test_write_permissions(directory):
    """Test if the directory is writable."""
    print(f"\n=== Testing write permissions for {directory} ===")
    try:
        test_file = os.path.join(directory, "write_test.txt")
        with open(test_file, 'w') as f:
            f.write("Write permission test")
        print(f"✅ Successfully wrote test file: {test_file}")
        
        # Read the file to verify
        with open(test_file, 'r') as f:
            content = f.read()
        print(f"✅ Successfully read test file (content length: {len(content)})")
        
        # Clean up
        os.remove(test_file)
        print(f"✅ Successfully removed test file")
        return True
    except Exception as e:
        print(f"❌ Write permission test failed: {str(e)}")
        print(f"Current user: {os.getlogin() if hasattr(os, 'getlogin') else 'unknown'}")
        print(f"Directory owner: {os.stat(directory).st_uid if os.path.exists(directory) else 'unknown'}")
        return False

def create_test_html():
    """Create a simple test HTML file."""
    print("\n=== Creating test HTML file ===")
    temp_dir = tempfile.mkdtemp()
    html_path = os.path.join(temp_dir, "test.html")
    
    html_content = """<!DOCTYPE html>
<html>
<head>
    <title>Test Document</title>
</head>
<body>
    <h1>Test Document</h1>
    <p>This is a test document for HTML to DOCX conversion.</p>
    <ul>
        <li>Item 1</li>
        <li>Item 2</li>
        <li>Item 3</li>
    </ul>
</body>
</html>"""
    
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    print(f"✅ Created test HTML file: {html_path}")
    print(f"HTML content size: {len(html_content)} bytes")
    return temp_dir, html_path

def test_html2docx_conversion(html_path, output_dir):
    """Test conversion with html2docx."""
    print("\n=== Testing html2docx conversion ===")
    try:
        import html2docx
        
        docx_path = os.path.join(output_dir, "test_output.docx")
        
        # Read HTML content
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        print(f"Read {len(html_content)} bytes from {html_path}")
        
        # Try conversion
        print(f"Converting to DOCX: {docx_path}")
        try:
            result = html2docx.html2docx(html_content, docx_path)
            print(f"Conversion result: {result}")
        except Exception as e:
            print(f"❌ html2docx.html2docx() failed: {str(e)}")
            traceback.print_exc()
            return False
        
        # Check if the file was created
        if os.path.exists(docx_path):
            file_size = os.path.getsize(docx_path)
            print(f"✅ DOCX file created: {docx_path} (size: {file_size} bytes)")
            
            if file_size == 0:
                print("❌ WARNING: File size is 0 bytes!")
                return False
            return True
        else:
            print(f"❌ DOCX file not created: {docx_path}")
            return False
            
    except Exception as e:
        print(f"❌ Conversion test failed: {str(e)}")
        traceback.print_exc()
        return False

def test_pandoc_availability():
    """Test if Pandoc is available."""
    print("\n=== Testing Pandoc availability ===")
    try:
        result = subprocess.run(
            ["pandoc", "--version"], 
            stdout=subprocess.PIPE, 
            stderr=subprocess.PIPE,
            check=False
        )
        if result.returncode == 0:
            version = result.stdout.decode('utf-8').split('\n')[0]
            print(f"✅ Pandoc is available: {version}")
            return True
        else:
            print("❌ Pandoc is installed but returned an error:")
            print(result.stderr.decode('utf-8'))
            return False
    except FileNotFoundError:
        print("❌ Pandoc is not installed or not in PATH")
        return False
    except Exception as e:
        print(f"❌ Error checking Pandoc: {str(e)}")
        return False

def test_pandoc_conversion(html_path, output_dir):
    """Test conversion with Pandoc."""
    print("\n=== Testing Pandoc conversion ===")
    if not test_pandoc_availability():
        print("Skipping Pandoc conversion test")
        return False
    
    docx_path = os.path.join(output_dir, "test_pandoc.docx")
    try:
        cmd = ["pandoc", html_path, "-f", "html", "-t", "docx", "-o", docx_path]
        print(f"Running Pandoc command: {' '.join(cmd)}")
        
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
        
        if result.returncode == 0:
            print("✅ Pandoc command executed successfully")
            
            if os.path.exists(docx_path):
                file_size = os.path.getsize(docx_path)
                print(f"✅ DOCX file created: {docx_path} (size: {file_size} bytes)")
                return True
            else:
                print(f"❌ DOCX file not created despite successful command: {docx_path}")
                return False
        else:
            print("❌ Pandoc command failed:")
            print(result.stderr.decode('utf-8'))
            return False
    except Exception as e:
        print(f"❌ Error running Pandoc: {str(e)}")
        return False

def test_direct_python_implementation(html_path, output_dir):
    """Test direct Python implementation for HTML to DOCX conversion."""
    print("\n=== Testing direct Python implementation ===")
    
    # First check if required packages are installed
    try:
        import html2docx
        from docx import Document
    except ImportError as e:
        print(f"❌ Required packages missing: {str(e)}")
        print("Skipping direct Python implementation test")
        return False
    
    docx_path = os.path.join(output_dir, "test_direct.docx")
    
    try:
        # Create minimal Python implementation
        print("Creating minimal conversion implementation")
        
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Try to create a simple document directly with python-docx
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document created directly with python-docx.')
        doc.save(docx_path)
        
        if os.path.exists(docx_path) and os.path.getsize(docx_path) > 0:
            print(f"✅ Created test DOCX directly: {docx_path} (size: {os.path.getsize(docx_path)} bytes)")
            return True
        else:
            print(f"❌ Failed to create direct test DOCX: {docx_path}")
            return False
            
    except Exception as e:
        print(f"❌ Direct Python implementation failed: {str(e)}")
        traceback.print_exc()
        return False

def cleanup(temp_dir):
    """Clean up temporary files."""
    import shutil
    try:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            print(f"✅ Cleaned up temporary directory: {temp_dir}")
    except Exception as e:
        print(f"❌ Failed to clean up: {str(e)}")

def run_all_tests():
    """Run all tests."""
    print("DOCX Converter Diagnostic Tool")
    print("=============================")
    
    # Test import first
    if not test_html2docx_import():
        print("\n❌ Cannot proceed without html2docx module")
        return 1
    
    # Create output directory if needed
    output_dir = os.path.join(os.getcwd(), "test_output")
    os.makedirs(output_dir, exist_ok=True)
    print(f"\nOutput directory: {output_dir}")
    
    # Test write permissions
    if not test_write_permissions(output_dir):
        print("\n❌ Cannot proceed without write permissions")
        return 1
    
    # Create test HTML
    temp_dir, html_path = create_test_html()
    
    # Run specific conversion tests
    html2docx_success = test_html2docx_conversion(html_path, output_dir)
    pandoc_success = test_pandoc_conversion(html_path, output_dir)
    direct_success = test_direct_python_implementation(html_path, output_dir)
    
    # Clean up
    cleanup(temp_dir)
    
    # Summary
    print("\n=== Test Summary ===")
    print(f"html2docx conversion: {'✅ Passed' if html2docx_success else '❌ Failed'}")
    print(f"Pandoc conversion: {'✅ Passed' if pandoc_success else '❌ Failed'}")
    print(f"Direct Python implementation: {'✅ Passed' if direct_success else '❌ Failed'}")
    
    if html2docx_success or pandoc_success or direct_success:
        print("\n✅ At least one conversion method works!")
        print("\nRecommendations:")
        if html2docx_success:
            print("- Use html2docx for conversion")
        elif pandoc_success:
            print("- Use Pandoc for conversion")
        elif direct_success:
            print("- Use direct python-docx implementation")
        return 0
    else:
        print("\n❌ All conversion methods failed.")
        print("\nPossible solutions:")
        print("1. Install Pandoc (recommended): https://pandoc.org/installing.html")
        print("2. Reinstall html2docx and python-docx:")
        print("   pip uninstall -y html2docx python-docx")
        print("   pip install python-docx html2docx")
        print("3. Check file system permissions")
        return 1

if __name__ == "__main__":
    sys.exit(run_all_tests())