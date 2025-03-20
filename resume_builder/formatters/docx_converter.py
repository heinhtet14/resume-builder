import os
import sys
import traceback
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxConverter:
    """Convert HTML to DOCX format using direct python-docx implementation."""
    
    def __init__(self):
        """Initialize the converter."""
        print("Initializing DocxConverter with direct python-docx implementation")
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are installed."""
        try:
            import docx
            import bs4
            print("Required dependencies are installed (python-docx, beautifulsoup4)")
        except ImportError as e:
            print(f"Missing dependencies: {e}")
            print("Attempting to install required packages...")
            
            try:
                import subprocess
                subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "beautifulsoup4"], 
                              check=True, capture_output=True)
                print("Successfully installed required packages")
            except Exception as install_err:
                print(f"Error installing packages: {str(install_err)}")
                print("Please manually install required packages:")
                print("pip install python-docx beautifulsoup4")
    
    def _direct_html_to_docx(self, html_content, output_path):
        """Convert HTML to DOCX using direct python-docx approach."""
        try:
            # Parse HTML content
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Create a new Document
            doc = Document()
            
            # Set document style
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Process HTML elements and add to document
            self._process_html_elements(soup.body, doc)
            
            # Save the document
            print(f"Saving document to: {output_path}")
            doc.save(output_path)
            
            # Verify file exists and has content
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"Successfully created DOCX: {output_path} (size: {os.path.getsize(output_path)} bytes)")
                return True
            else:
                print(f"File not created or empty: {output_path}")
                return False
                
        except Exception as e:
            print(f"Error converting HTML to DOCX: {str(e)}")
            traceback.print_exc()
            return False
    
    def _process_html_elements(self, parent, doc, paragraph=None):
        """Process HTML elements recursively and add to document."""
        if not parent:
            return
        
        for element in parent.children:
            if element.name is None:  # Text node
                if element.string and element.string.strip():
                    if paragraph is None:
                        paragraph = doc.add_paragraph()
                    run = paragraph.add_run(element.string.strip())
            elif element.name == 'h1':
                paragraph = doc.add_heading(level=1)
                run = paragraph.add_run(element.get_text().strip())
                paragraph = None
            elif element.name == 'h2':
                paragraph = doc.add_heading(level=2)
                run = paragraph.add_run(element.get_text().strip())
                paragraph = None
            elif element.name == 'h3':
                paragraph = doc.add_heading(level=3)
                run = paragraph.add_run(element.get_text().strip())
                paragraph = None
            elif element.name == 'p':
                paragraph = doc.add_paragraph()
                # Apply text formatting for the paragraph
                if 'class' in element.attrs:
                    if 'text-center' in element['class']:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Process children of the paragraph
                self._process_html_elements(element, doc, paragraph)
                paragraph = None
            elif element.name == 'strong' or element.name == 'b':
                if paragraph is None:
                    paragraph = doc.add_paragraph()
                run = paragraph.add_run(element.get_text().strip())
                run.bold = True
            elif element.name == 'em' or element.name == 'i':
                if paragraph is None:
                    paragraph = doc.add_paragraph()
                run = paragraph.add_run(element.get_text().strip())
                run.italic = True
            elif element.name == 'ul':
                # Process list items
                for li in element.find_all('li', recursive=False):
                    paragraph = doc.add_paragraph(style='List Bullet')
                    run = paragraph.add_run(li.get_text().strip())
                paragraph = None
            elif element.name == 'ol':
                # Process ordered list items
                for li in element.find_all('li', recursive=False):
                    paragraph = doc.add_paragraph(style='List Number')
                    run = paragraph.add_run(li.get_text().strip())
                paragraph = None
            elif element.name == 'a':
                if paragraph is None:
                    paragraph = doc.add_paragraph()
                run = paragraph.add_run(element.get_text().strip())
                # TODO: Add hyperlink if needed
            elif element.name == 'br':
                if paragraph is None:
                    paragraph = doc.add_paragraph()
                run = paragraph.add_run('\n')
            elif element.name == 'div':
                # Process div contents
                self._process_html_elements(element, doc, paragraph)
            elif element.name == 'span':
                if paragraph is None:
                    paragraph = doc.add_paragraph()
                run = paragraph.add_run(element.get_text().strip())
            elif element.name in ['table', 'tr', 'td', 'th']:
                # Skip table processing for now (complex)
                pass
            else:
                # Process other elements recursively
                self._process_html_elements(element, doc, paragraph)
    
    def convert_html_to_docx(self, html_content=None, html_file=None, output_path=None):
        """
        Convert HTML to DOCX format.
        
        Args:
            html_content: HTML content as a string
            html_file: Path to an HTML file
            output_path: Path to save the DOCX
            
        Returns:
            The path to the generated DOCX file, or None if conversion failed
        """
        if html_content is None and html_file is None:
            raise ValueError("Either html_content or html_file must be provided")
        
        if output_path is None:
            if html_file:
                output_path = os.path.splitext(html_file)[0] + ".docx"
            else:
                output_path = "resume.docx"
        
        print(f"Converting HTML to DOCX: {output_path}")
        
        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        
        try:
            # Get HTML content
            if html_file:
                print(f"Reading HTML from file: {html_file}")
                with open(html_file, 'r', encoding='utf-8') as f:
                    html_content = f.read()
            
            # Validate HTML content
            if not html_content or len(html_content.strip()) == 0:
                print("ERROR: HTML content is empty")
                return None
            
            print(f"HTML content length: {len(html_content)} bytes")
            
            # Try direct conversion
            print("Converting HTML to DOCX using direct python-docx implementation...")
            success = self._direct_html_to_docx(html_content, output_path)
            
            if success:
                print(f"Successfully converted HTML to DOCX: {output_path}")
                return output_path
            else:
                print(f"Failed to convert HTML to DOCX")
                
                # Fall back to creating a basic DOCX with resume content
                try:
                    print("Attempting fallback method (creating a basic DOCX)...")
                    doc = Document()
                    
                    # Extract basic text from HTML using BeautifulSoup
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    # Add title (if available)
                    title = soup.title.string if soup.title else "Resume"
                    doc.add_heading(title, 0)
                    
                    # Extract text from all paragraphs
                    for p in soup.find_all(['p', 'div']):
                        text = p.get_text().strip()
                        if text:
                            doc.add_paragraph(text)
                    
                    # Save the basic document
                    doc.save(output_path)
                    print(f"Created basic DOCX with text content: {output_path}")
                    return output_path
                    
                except Exception as fallback_err:
                    print(f"Fallback method failed: {str(fallback_err)}")
                    return None
                    
        except Exception as e:
            print(f"Error in HTML to DOCX conversion: {str(e)}")
            traceback.print_exc()
            return None