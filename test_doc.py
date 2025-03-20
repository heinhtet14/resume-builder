#!/usr/bin/env python3
# convert_resume.py - Script to convert HTML resume to DOCX format

import os
import sys
from pathlib import Path

def convert_resume():
    """Convert the HTML resume to DOCX format."""
    print("Resume HTML to DOCX Converter")
    print("=============================")
    
    # Define the path to the HTML resume file
    output_dir = Path("output")
    html_file = output_dir / "resume.html"
    docx_file = output_dir / "resume.docx"
    
    # Check if HTML file exists
    if not html_file.exists():
        print(f"ERROR: HTML file not found: {html_file}")
        return 1
    
    print(f"HTML file found: {html_file} (size: {html_file.stat().st_size} bytes)")
    
    # Install necessary packages if they aren't already installed
    try:
        import docx
        import bs4
    except ImportError:
        print("Installing required packages...")
        import subprocess
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "beautifulsoup4"], 
                       check=True)
    
    # Create a simple document
    from docx import Document
    from bs4 import BeautifulSoup
    
    print("Reading HTML content...")
    with open(html_file, "r", encoding="utf-8") as f:
        html_content = f.read()
    
    print(f"Read {len(html_content)} bytes of HTML")
    
    print("Creating DOCX document...")
    doc = Document()
    
    soup = BeautifulSoup(html_content, "html.parser")
    
    # Add name as title
    if soup.h1:
        doc.add_heading(soup.h1.get_text().strip(), 0)
    
    # Add contact info
    if soup.find(class_="contact-info"):
        contact_text = soup.find(class_="contact-info").get_text().strip()
        doc.add_paragraph(contact_text)
    
    # Add summary
    if soup.find(class_="summary"):
        summary_text = soup.find(class_="summary").get_text().strip()
        doc.add_paragraph(summary_text)
    
    # Add sections
    for section in soup.find_all("section"):
        section_title = section.find("h2")
        if section_title:
            doc.add_heading(section_title.get_text().strip(), 1)
        
        # Process jobs, education, projects
        for item in section.find_all(["div"], class_=["job", "education", "project"]):
            # Add title/header
            header = item.find(class_=["job-header", "edu-header", "project-header"])
            if header:
                p = doc.add_paragraph()
                title_part = header.find(class_=["job-title", "degree", "project-name"])
                if title_part:
                    p.add_run(title_part.get_text().strip()).bold = True
                
                # Add duration
                duration = header.find(class_=["job-duration", "edu-year", "project-duration"])
                if duration:
                    p.add_run(" - " + duration.get_text().strip())
            
            # Add responsibilities
            for ul in item.find_all("ul"):
                for li in ul.find_all("li"):
                    doc.add_paragraph(li.get_text().strip(), style="List Bullet")
    
    print(f"Saving DOCX to: {docx_file}")
    doc.save(docx_file)
    
    if docx_file.exists():
        print(f"✅ Successfully created DOCX: {docx_file} (size: {docx_file.stat().st_size} bytes)")
        return 0
    else:
        print(f"❌ Failed to create DOCX file: {docx_file}")
        return 1

if __name__ == "__main__":
    sys.exit(convert_resume())